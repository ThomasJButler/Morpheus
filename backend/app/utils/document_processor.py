"""
Document processing utilities for various file formats.
Supports PDF, TXT, MD, and DOCX files.
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Handles extraction of text from various document formats.
    """

    def __init__(self):
        """Initialize document processor."""
        self.supported_types = settings.supported_file_types_list
        logger.info(f"DocumentProcessor initialized. Supported types: {self.supported_types}")

    def process_file(self, file_path: str) -> Dict:
        """
        Process a file and extract text content.

        Args:
            file_path: Path to the file

        Returns:
            Dict with extracted text and metadata

        Raises:
            ValueError: If file type is not supported
        """
        path = Path(file_path)
        file_extension = path.suffix.lower().lstrip(".")

        if file_extension not in self.supported_types:
            raise ValueError(
                f"Unsupported file type: {file_extension}. "
                f"Supported types: {', '.join(self.supported_types)}"
            )

        logger.info(f"Processing file: {path.name} ({file_extension})")

        # Route to appropriate processor
        if file_extension == "pdf":
            return self._process_pdf(file_path)
        elif file_extension == "txt":
            return self._process_txt(file_path)
        elif file_extension == "md":
            return self._process_markdown(file_path)
        elif file_extension == "docx":
            return self._process_docx(file_path)
        else:
            raise ValueError(f"No processor implemented for {file_extension}")

    def _process_pdf(self, file_path: str) -> Dict:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Dict with text and metadata
        """
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install pypdf2")

        path = Path(file_path)
        text_content = []
        metadata = {
            "source": path.name,
            "file_type": "pdf",
            "pages": 0,
        }

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append({
                            "text": text,
                            "page": page_num,
                            "source": path.name,
                        })

            logger.info(f"Extracted text from {metadata['pages']} pages in {path.name}")

            return {
                "content": text_content,
                "metadata": metadata,
                "success": True,
            }

        except Exception as e:
            logger.error(f"Error processing PDF {path.name}: {e}")
            return {
                "content": [],
                "metadata": metadata,
                "success": False,
                "error": str(e),
            }

    def _process_txt(self, file_path: str) -> Dict:
        """
        Extract text from TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            Dict with text and metadata
        """
        path = Path(file_path)
        metadata = {
            "source": path.name,
            "file_type": "txt",
        }

        try:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()

            return {
                "content": [{
                    "text": text,
                    "source": path.name,
                }],
                "metadata": metadata,
                "success": True,
            }

        except Exception as e:
            logger.error(f"Error processing TXT {path.name}: {e}")
            return {
                "content": [],
                "metadata": metadata,
                "success": False,
                "error": str(e),
            }

    def _process_markdown(self, file_path: str) -> Dict:
        """
        Extract text from Markdown file.

        Args:
            file_path: Path to MD file

        Returns:
            Dict with text and metadata
        """
        # For now, treat markdown like plain text
        # Could enhance with markdown parsing library
        return self._process_txt(file_path)

    def _process_docx(self, file_path: str) -> Dict:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dict with text and metadata
        """
        if DocxDocument is None:
            raise ImportError(
                "python-docx is required for DOCX processing. Install with: pip install python-docx"
            )

        path = Path(file_path)
        text_content = []
        metadata = {
            "source": path.name,
            "file_type": "docx",
            "paragraphs": 0,
            "tables": 0,
        }

        try:
            # Load DOCX document
            doc = DocxDocument(file_path)

            # Process paragraphs
            current_text = []
            paragraph_count = 0

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    current_text.append(text)
                    paragraph_count += 1

                    # Group paragraphs into chunks (every 10 paragraphs or on heading).
                    # `paragraph.style` can be None for some .docx files
                    # (Word omits the style ref when the paragraph inherits from
                    # the document default), so null-check before reading .name.
                    if paragraph_count % 10 == 0 or (
                        paragraph.style is not None
                        and paragraph.style.name
                        and paragraph.style.name.startswith("Heading")
                    ):
                        if current_text:
                            text_content.append({
                                "text": "\n".join(current_text),
                                "source": path.name,
                                "paragraph_index": paragraph_count,
                            })
                            current_text = []

            # Add any remaining text
            if current_text:
                text_content.append({
                    "text": "\n".join(current_text),
                    "source": path.name,
                    "paragraph_index": paragraph_count,
                })

            metadata["paragraphs"] = paragraph_count

            # Process tables
            table_count = 0
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))

                if table_text:
                    table_count += 1
                    text_content.append({
                        "text": "\n".join(table_text),
                        "source": path.name,
                        "table_index": table_count,
                    })

            metadata["tables"] = table_count

            logger.info(
                f"Extracted text from {paragraph_count} paragraphs "
                f"and {table_count} tables in {path.name}"
            )

            return {
                "content": text_content,
                "metadata": metadata,
                "success": True,
            }

        except Exception as e:
            logger.error(f"Error processing DOCX {path.name}: {e}")
            return {
                "content": [],
                "metadata": metadata,
                "success": False,
                "error": str(e),
            }

    def validate_file_size(self, file_path: str) -> bool:
        """
        Check if file size is within limits.

        Args:
            file_path: Path to file

        Returns:
            bool: True if file is within size limit
        """
        path = Path(file_path)
        size_mb = path.stat().st_size / (1024 * 1024)

        if size_mb > settings.max_file_size_mb:
            logger.warning(
                f"File {path.name} ({size_mb:.2f}MB) exceeds limit ({settings.max_file_size_mb}MB)"
            )
            return False

        return True


# Convenience function
def process_document(file_path: str) -> Dict:
    """
    Quick document processing function.

    Args:
        file_path: Path to document

    Returns:
        Processed document dict
    """
    processor = DocumentProcessor()
    return processor.process_file(file_path)
